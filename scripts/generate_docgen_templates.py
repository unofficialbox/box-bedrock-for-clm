#!/usr/bin/env python3
"""Generate Box DocGen-ready CLM Word templates with deterministic styling."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "docgen"
BLUE = "2E74B5"
DARK = "1F4D78"
GRAY = "5F6368"
LIGHT = "F2F4F7"


def set_font(run, size=11, bold=False, color="000000", italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    total = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def configure(doc, running_label, footer_label="Acme Technologies | CLM-2026-Northstar"):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 12, 8, 4, DARK),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_font(header.add_run(running_label), 9, True, GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run(footer_label), 8, False, GRAY)


def add_title(doc, kicker, title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(kicker.upper()), 9, True, BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(title), 23, True, "000000")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    set_font(p.add_run(subtitle), 12, False, GRAY)


def add_key_values(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cells[0], LIGHT)
        shade(cells[1], "FFFFFF")
        set_font(cells[0].paragraphs[0].add_run(label), 10, True, DARK)
        set_font(cells[1].paragraphs[0].add_run(value), 10.5)
    set_table_geometry(table, [2700, 6660])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_section_text(doc, heading, value):
    doc.add_heading(heading, level=2)
    p = doc.add_paragraph()
    set_font(p.add_run(value), 11)


def approval_memo():
    doc = Document()
    configure(doc, "CLM Approval Memo")
    add_title(doc, "Decision required", "Contract Approval Memo", "Northstar Health System agreement package")
    add_key_values(doc, [
        ("Contract ID", "{{contract.id}}"),
        ("Counterparty", "{{contract.counterparty}}"),
        ("Contract type", "{{contract.type}}"),
        ("Deal value", "{{contract.currency}}{{contract.dealValue}}"),
        ("Region / data", "{{contract.region}} / {{contract.dataCategory}}"),
        ("Target signature", "{{contract.targetSignatureDate}}"),
        ("Business owner", "{{contract.businessOwner}}"),
    ])
    add_section_text(doc, "Executive summary", "{{approval.executiveSummary}}")
    add_key_values(doc, [
        ("Legal review", "{{reviews.legalStatus}} - {{reviews.legalSummary}}"),
        ("Finance review", "{{reviews.financeStatus}} - {{reviews.financeSummary}}"),
        ("Privacy / security", "{{reviews.privacySecurityStatus}} - {{reviews.privacySecuritySummary}}"),
        ("Overall risk", "{{approval.riskLevel}}"),
    ])
    add_section_text(doc, "Recommendation", "{{approval.recommendation}}")
    add_key_values(doc, [
        ("Approver", "{{approval.approver}}"),
        ("Decision", "{{approval.decision}}"),
        ("Decision date", "{{approval.decisionDate}}"),
    ])
    return doc


def order_summary():
    doc = Document()
    configure(doc, "Commercial Order Summary")
    add_title(doc, "Commercial record", "Order Form Summary", "Structured deal terms for review and approval")
    add_key_values(doc, [
        ("Contract ID", "{{contract.id}}"),
        ("Customer", "{{contract.counterparty}}"),
        ("Annual value", "{{commercial.currency}}{{commercial.annualValue}}"),
        ("Total value", "{{commercial.currency}}{{commercial.totalValue}}"),
        ("Term", "{{commercial.termMonths}} months"),
        ("Effective date", "{{commercial.effectiveDate}}"),
        ("Payment terms", "{{commercial.paymentTerms}}"),
        ("Renewal", "{{commercial.renewalType}}"),
        ("Renewal notice", "{{commercial.renewalNoticeDays}} days"),
        ("Governing law", "{{commercial.governingLaw}}"),
    ])
    add_section_text(doc, "Commercial exception", "{{commercial.exceptionSummary}}")
    add_section_text(doc, "Approval note", "{{commercial.approvalNote}}")
    return doc


def renewal_notice():
    doc = Document()
    configure(doc, "Contract Renewal Notice")
    add_title(doc, "Contract notice", "Renewal Notice", "Formal notice generated from tracked obligations")
    add_key_values(doc, [
        ("Notice date", "{{notice.date}}"),
        ("Recipient", "{{recipient.name}}, {{recipient.title}}"),
        ("Recipient company", "{{recipient.company}}"),
        ("Recipient email", "{{recipient.email}}"),
        ("Contract ID", "{{contract.id}}"),
        ("Agreement", "{{contract.name}}"),
        ("Current term ends", "{{renewal.currentTermEndDate}}"),
        ("Notice deadline", "{{renewal.noticeDeadline}}"),
    ])
    add_section_text(doc, "Notice", "{{renewal.noticeText}}")
    add_section_text(doc, "Requested action", "{{renewal.requestedAction}}")
    add_key_values(doc, [
        ("Sender", "{{sender.name}}, {{sender.title}}"),
        ("Sender company", "{{sender.company}}"),
        ("Sender email", "{{sender.email}}"),
    ])
    return doc


def msa_redline():
    """Current-year MSA still in redline: a Word contract with visible redline markup (not a merge template)."""
    doc = Document()
    configure(doc, "CLM MSA Redline 2026")
    add_title(
        doc,
        "In negotiation",
        "Master Services Agreement - Redline",
        "Acme Technologies, Inc. and Northstar Health System | CLM-2026-Northstar | Draft, in redline",
    )

    def clause(heading, body, redline):
        doc.add_heading(heading, level=2)
        set_font(doc.add_paragraph().add_run(body), 11)
        rp = doc.add_paragraph()
        set_font(rp.add_run("NORTHSTAR REDLINE: " + redline), 11, bold=True, color="B91C1C")

    clause(
        "1. Fees, invoicing, and taxes",
        "Customer will pay the fees stated in each Order Form. Subscription fees are invoiced annually in advance; undisputed amounts are due forty-five days from invoice date.",
        "Customer requests Net 90 payment terms and a unilateral right to offset alleged damages against invoiced amounts.",
    )
    clause(
        "2. Service levels and support",
        "Service credits are Customer's sole monetary remedy for a service-level failure and will not exceed twenty percent of the monthly subscription fees for the affected Service.",
        "Customer proposes uncapped service credits, a full monthly refund below 99.5 percent availability, and termination rights after any two failures in a rolling six-month period.",
    )
    clause(
        "3. Information security",
        "Acme will notify Customer of a confirmed Security Incident without undue delay and no later than seventy-two hours after confirmation. Annual audits are permitted with reasonable notice.",
        "Customer requests notice within twelve hours of any suspected event, participation in Acme's forensic investigation, and unlimited on-site audits.",
    )
    clause(
        "4. Limitation of liability",
        "Except for Excluded Claims, each Party's aggregate liability will not exceed the fees paid or payable under the affected Order Form during the preceding twelve months.",
        "Customer deletes the liability cap and proposes unlimited liability for direct and indirect damages, including lost revenue and regulatory penalties.",
    )
    clause(
        "5. Term and termination",
        "Each Order Form renews for successive one-year periods unless either Party provides at least ninety days' written notice before the current term ends.",
        "Customer adds termination for convenience on thirty days' notice with a pro rata refund of prepaid fees.",
    )
    clause(
        "6. Governing law",
        "This Agreement is governed by Delaware law; courts in Wilmington, Delaware have exclusive jurisdiction.",
        "Customer replaces Delaware law and venue with Minnesota law and exclusive venue in Hennepin County, Minnesota.",
    )
    doc.add_heading("Open items", level=2)
    set_font(
        doc.add_paragraph().add_run(
            "Red text marks the counterparty's proposed changes still under negotiation. This document is a draft in Word and is not executed."
        ),
        10,
        italic=True,
        color=GRAY,
    )
    return doc


def counter_position():
    """The document Acme sends back when the deal is on the customer's paper.

    Generating "the agreement" is the wrong artifact here -- the agreement is Calder's, and
    Acme is not going to redraft it. What Acme produces is a position: which clause is out of
    policy, what the approved fallback text is, who owns the deviation, and what the
    counterparty already agreed to the last two times. Every one of those already exists as
    governed content -- the clause library Hub supplies the fallback and its owner, the prior
    executed agreements supply the precedent -- so this template is the place they are
    assembled, not the place they are invented.
    """
    doc = Document()
    configure(doc, "CLM Counter-Position", "Acme Technologies | {{contract.id}}")
    add_title(
        doc,
        "Negotiation position",
        "Counter-Position Memo",
        "Prepared from the approved clause library and prior executed agreements",
    )
    add_key_values(doc, [
        ("Contract ID", "{{contract.id}}"),
        ("Counterparty", "{{contract.counterparty}}"),
        ("Their paper", "{{contract.paperReference}}"),
        ("Deal value", "{{contract.dealValue}}"),
        ("Status", "{{contract.status}}"),
        ("Prepared", "{{memo.preparedOn}}"),
    ])
    add_section_text(doc, "The clause at issue", "{{position.clauseAtIssue}}")
    add_section_text(doc, "What their draft says", "{{position.theirPosition}}")
    add_key_values(doc, [
        ("Approved position", "{{position.approvedClauseId}} - {{position.approvedPosition}}"),
        ("Approved fallback", "{{position.fallbackClauseId}} - {{position.fallbackPosition}}"),
        ("Deviation owner", "{{position.owner}}"),
        ("Risk", "{{position.risk}}"),
    ])
    add_section_text(doc, "What they agreed before", "{{precedent.summary}}")
    add_section_text(doc, "Proposed counter-position", "{{position.counterProposal}}")
    add_key_values(doc, [
        ("Prepared by", "{{memo.preparedBy}}"),
        ("Requires approval from", "{{position.owner}}"),
    ])
    add_section_text(
        doc,
        "Note",
        "This memo is a draft prepared for review. It is not an approval, and it does not "
        "authorise signature. {{position.owner}} owns the decision on this deviation.",
    )
    return doc


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    templates = {
        "clm-approval-memo-template.docx": approval_memo(),
        "clm-counter-position-template.docx": counter_position(),
        "clm-order-summary-template.docx": order_summary(),
        "clm-renewal-notice-template.docx": renewal_notice(),
        "northstar-msa-2026-redline.docx": msa_redline(),
    }
    for name, document in templates.items():
        document.core_properties.title = name.removesuffix(".docx").replace("-", " ").title()
        document.core_properties.subject = "Box DocGen template for the Northstar CLM demo"
        document.core_properties.author = "Acme Technologies CLM Demo"
        document.save(OUTPUT / name)
        print(OUTPUT / name)


if __name__ == "__main__":
    main()
